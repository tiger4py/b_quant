"""
Markdown → Word (.docx) 转换器
使用 mistune 解析 + python-docx 生成，支持标题、表格、引用、代码块、列表等
"""
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import mistune

# ---------- 样式配置 ----------
FONT_NAME = 'Microsoft YaHei'
FONT_SIZE_BODY = Pt(10.5)
FONT_SIZE_H1 = Pt(18)
FONT_SIZE_H2 = Pt(15)
FONT_SIZE_H3 = Pt(13)
FONT_SIZE_H4 = Pt(11)
FONT_SIZE_CODE = Pt(9)
LINE_SPACING = 1.35


def set_cell_shading(cell, color_hex):
    """设置单元格底色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '999999'))
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_run_with_font(paragraph, text, bold=False, italic=False, font_size=None, color=None, font_name=None):
    """添加带格式的文本段"""
    run = paragraph.add_run(text)
    run.font.name = font_name or FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name or FONT_NAME)
    run.font.size = font_size or FONT_SIZE_BODY
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run


def parse_inline_text(paragraph, text):
    """解析行内 markdown 并添加到段落"""
    # 处理 **bold**, *italic*, `code`
    pattern = re.compile(
        r'(\*\*(.+?)\*\*)|'          # bold
        r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)|'  # italic
        r'`([^`]+)`'                  # code
    )
    last_end = 0
    for m in pattern.finditer(text):
        # 前面的普通文本
        if m.start() > last_end:
            add_run_with_font(paragraph, text[last_end:m.start()])
        if m.group(1):  # bold
            add_run_with_font(paragraph, m.group(2), bold=True)
        elif m.group(3):  # italic
            add_run_with_font(paragraph, m.group(3), italic=True)
        elif m.group(4):  # code
            add_run_with_font(paragraph, m.group(4), font_size=FONT_SIZE_CODE, color=(180, 60, 60))
        last_end = m.end()
    # 剩余文本
    if last_end < len(text):
        add_run_with_font(paragraph, text[last_end:])


class DocxRenderer(mistune.HTMLRenderer):
    """将 mistune AST 渲染为 python-docx 文档"""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None

    def _add_heading(self, text, level):
        p = self.doc.add_paragraph()
        size_map = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3, 4: FONT_SIZE_H4}
        add_run_with_font(p, text, bold=True, font_size=size_map.get(level, FONT_SIZE_BODY))
        p.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = LINE_SPACING
        return p

    def heading(self, text, level, **attrs):
        return self._add_heading(text, level)

    def paragraph(self, text):
        # 跳过空段落和纯分隔符
        if not text or text.strip() in ('---', '***', '___'):
            return ''
        p = self.doc.add_paragraph()
        parse_inline_text(p, text)
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(4)
        return p

    def block_quote(self, text):
        p = self.doc.add_paragraph()
        add_run_with_font(p, text, italic=True, color=(100, 100, 100))
        # 左边缩进
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.line_spacing = LINE_SPACING
        return p

    def block_code(self, code, info=None):
        p = self.doc.add_paragraph()
        add_run_with_font(p, code, font_size=FONT_SIZE_CODE, color=(60, 60, 60))
        # 灰色背景模拟
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.line_spacing = 1.0
        return p

    def codespan(self, text):
        return text  # 由 parse_inline_text 处理

    def thematic_break(self):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        # 添加下划线作为分隔
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'CCCCCC')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return ''

    def list_item(self, text, ordered=False, level=0, **attrs):
        p = self.doc.add_paragraph()
        prefix = '• ' if not ordered else f'{ordered}. '
        add_run_with_font(p, prefix, bold=True)
        parse_inline_text(p, text)
        p.paragraph_format.left_indent = Cm(0.8 + level * 0.6)
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(2)
        return p

    def table(self, text):
        """由 mistune 插件的表格解析调用 — 这里用自定义 table 处理"""
        return text

    def table_row(self, text):
        return text

    def table_cell(self, text, align=None, is_head=False):
        return text

    def strong(self, text):
        return text

    def emphasis(self, text):
        return text

    def inline_html(self, text):
        return text

    def linebreak(self):
        return '\n'

    def text(self, text):
        return text

    def blank_line(self):
        return ''

    def link(self, text, url, title=None):
        return text

    def image(self, src, alt='', title=None):
        return alt or src

    def finalize(self, data):
        return data


def build_docx_from_md(md_path, docx_path):
    """主转换函数：逐块解析 markdown，手动处理表格和复杂结构"""
    md_text = Path(md_path).read_text(encoding='utf-8')

    doc = Document()

    # ---------- 页面设置 ----------
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # ---------- 设置默认字体 ----------
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE_BODY
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    style.paragraph_format.line_spacing = LINE_SPACING

    lines = md_text.split('\n')

    i = 0
    in_table = False
    table_rows = []
    in_code_block = False
    code_lines = []
    in_blockquote = False
    quote_lines = []

    def flush_paragraph():
        """非表格、非代码块上下文的刷新不需要，段落逐行处理"""
        pass

    def add_table_to_doc(rows):
        """将 markdown 表格行转换为 docx 表格"""
        if not rows:
            return
        # 解析 markdown 表格行
        parsed = []
        col_count = 0
        for row in rows:
            if re.match(r'^[\s|:\-]+$', row):  # 分隔行
                continue
            cells = [c.strip() for c in row.strip().strip('|').split('|')]
            parsed.append(cells)
            col_count = max(col_count, len(cells))

        if not parsed:
            return

        table = doc.add_table(rows=len(parsed), cols=col_count)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row_cells in enumerate(parsed):
            for ci, cell_text in enumerate(row_cells):
                if ci >= col_count:
                    break
                cell = table.cell(ri, ci)
                # 清除默认段落
                cell.paragraphs[0].clear()
                p = cell.paragraphs[0]
                if ri == 0:
                    # 表头：深色背景 + 白色粗体
                    add_run_with_font(p, cell_text, bold=True, font_size=Pt(9.5), color=(255, 255, 255))
                    set_cell_shading(cell, '2F5496')
                else:
                    add_run_with_font(p, cell_text, font_size=Pt(9.5))
                    if ri % 2 == 0:
                        set_cell_shading(cell, 'F2F2F2')
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)

        # 表后空行
        doc.add_paragraph()

    def add_code_block_to_doc(code_lines):
        if not code_lines:
            return
        # 去掉首尾空行
        while code_lines and not code_lines[0].strip():
            code_lines.pop(0)
        while code_lines and not code_lines[-1].strip():
            code_lines.pop()
        for cl in code_lines:
            p = doc.add_paragraph()
            add_run_with_font(p, cl, font_size=FONT_SIZE_CODE, color=(60, 60, 60), font_name='Consolas')
            p.paragraph_format.left_indent = Cm(0.6)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)

    def add_quote_to_doc(quote_lines):
        if not quote_lines:
            return
        text = ' '.join(quote_lines)
        p = doc.add_paragraph()
        add_run_with_font(p, text, italic=True, color=(80, 80, 80))
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(4)

    def is_table_separator(line):
        return bool(re.match(r'^[\s|:\-]+$', line)) and '---' in line

    def is_table_row(line):
        return line.strip().startswith('|') and line.strip().endswith('|')

    while i < len(lines):
        line = lines[i]

        # ---- 代码块 ----
        if line.strip().startswith('```'):
            if in_code_block:
                add_code_block_to_doc(code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ---- 表格 ----
        if is_table_row(line):
            if not in_table:
                # 看看下一行是不是分隔行
                if i + 1 < len(lines) and is_table_separator(lines[i + 1]):
                    in_table = True
                    table_rows = [line]
                    i += 1
                    continue
            else:
                table_rows.append(line)
                # 下一行不是表格行 → 结束
                if i + 1 >= len(lines) or not is_table_row(lines[i + 1]):
                    add_table_to_doc(table_rows)
                    table_rows = []
                    in_table = False
                i += 1
                continue
        elif in_table and is_table_separator(line):
            table_rows.append(line)
            i += 1
            continue

        # ---- 引用块 ----
        if line.strip().startswith('> '):
            quote_lines.append(line.strip()[2:])
            in_blockquote = True
            i += 1
            continue
        elif in_blockquote and line.strip():
            # 非空行但不再以 > 开头 → 可能引用结束，也可能内部继续
            if line.strip().startswith('>'):
                quote_lines.append(line.strip()[2:])
                i += 1
                continue
            else:
                add_quote_to_doc(quote_lines)
                quote_lines = []
                in_blockquote = False
                # 继续处理当前行
        elif in_blockquote:
            add_quote_to_doc(quote_lines)
            quote_lines = []
            in_blockquote = False

        # ---- 水平分割线 ----
        if line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '4')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ---- 标题 ----
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            p = doc.add_paragraph()
            size_map = {1: FONT_SIZE_H1, 2: FONT_SIZE_H2, 3: FONT_SIZE_H3, 4: FONT_SIZE_H4,
                        5: Pt(10.5), 6: Pt(10)}
            add_run_with_font(p, text, bold=True, font_size=size_map.get(level, FONT_SIZE_BODY))
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = LINE_SPACING
            i += 1
            continue

        # ---- 无序列表 ----
        list_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if list_match:
            indent_level = len(list_match.group(1)) // 2
            text = list_match.group(2)
            p = doc.add_paragraph()
            add_run_with_font(p, '• ', bold=True)
            parse_inline_text(p, text)
            p.paragraph_format.left_indent = Cm(0.8 + indent_level * 0.6)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ---- 有序列表 ----
        ol_match = re.match(r'^(\s*)\d+[.)]\s+(.+)$', line)
        if ol_match:
            indent_level = len(ol_match.group(1)) // 2
            text = ol_match.group(2)
            num = re.match(r'\d+', line.strip())
            prefix = (num.group(0) if num else '1') + '. '
            p = doc.add_paragraph()
            add_run_with_font(p, prefix, bold=True)
            parse_inline_text(p, text)
            p.paragraph_format.left_indent = Cm(0.8 + indent_level * 0.6)
            p.paragraph_format.line_spacing = LINE_SPACING
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ---- 空行 ----
        if not line.strip():
            i += 1
            continue

        # ---- 普通段落 ----
        p = doc.add_paragraph()
        parse_inline_text(p, line)
        p.paragraph_format.line_spacing = LINE_SPACING
        p.paragraph_format.space_after = Pt(4)
        i += 1

    # 处理文件末尾未闭合的元素
    if in_code_block:
        add_code_block_to_doc(code_lines)
    if in_table and table_rows:
        add_table_to_doc(table_rows)
    if in_blockquote and quote_lines:
        add_quote_to_doc(quote_lines)

    # ---------- 保存 ----------
    doc.save(docx_path)
    print(f'[OK] 已生成: {docx_path}')


if __name__ == '__main__':
    md_file = sys.argv[1] if len(sys.argv) > 1 else '八字命理分析_乙木日主.md'
    docx_file = Path(md_file).with_suffix('.docx')
    build_docx_from_md(md_file, docx_file)
